from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_sample_legal_document():
    """Create a sample legal document for testing"""
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph("EMPLOYMENT AGREEMENT")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()  # Empty line
    
    # Introduction
    intro = doc.add_paragraph(
        "This Employment Agreement (the 'Agreement') is entered into on January 1, 2026, "
        "between Acme Corporation ('Employer') and John Doe ('Employee')."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Section 1
    doc.add_paragraph()
    section1_title = doc.add_paragraph("1. Position and Duties")
    section1_title.runs[0].bold = True
    
    section1_content = doc.add_paragraph(
        "The Employee shall serve as Senior Software Engineer and shall perform such duties "
        "as are customarily associated with such position. The Employee shall report directly "
        "to the Chief Technology Officer."
    )
    section1_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Section 2
    doc.add_paragraph()
    section2_title = doc.add_paragraph("2. Compensation")
    section2_title.runs[0].bold = True
    
    section2_content = doc.add_paragraph(
        "The Employee shall receive an annual base salary of $120,000, payable in accordance "
        "with the Employer's standard payroll practices. The Employee shall be eligible for "
        "an annual performance bonus of up to 15% of base salary."
    )
    section2_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Section 3
    doc.add_paragraph()
    section3_title = doc.add_paragraph("3. Confidentiality")
    section3_title.runs[0].bold = True
    
    section3_content = doc.add_paragraph(
        "The Employee agrees to maintain the confidentiality of all proprietary information "
        "and trade secrets of the Employer during and after the term of employment."
    )
    section3_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig1 = doc.add_paragraph("_________________________")
    sig1.add_run("\nEmployer Signature")
    
    doc.add_paragraph()
    
    sig2 = doc.add_paragraph("_________________________")
    sig2.add_run("\nEmployee Signature")
    
    # Save document
    output_dir = "tests/fixtures"
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, "sample_employment_agreement.docx")
    doc.save(output_path)
    
    print(f"✅ Sample document created: {output_path}")
    return output_path


if __name__ == "__main__":
    create_sample_legal_document()
