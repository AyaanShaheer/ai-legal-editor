from docx import Document as DocxDocument
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional
from io import BytesIO
from app.core.logging import logger
from app.core.exceptions import DocumentProcessingException


class DOCXParser:
    """Parse DOCX files and extract structured content"""
    
    def __init__(self, file_content: bytes):
        """
        Initialize parser with file content
        
        Args:
            file_content: Raw bytes of DOCX file
        """
        try:
            self.document = DocxDocument(BytesIO(file_content))
            logger.info("DOCX document loaded successfully")
        except Exception as e:
            logger.error(f"Failed to load DOCX: {str(e)}")
            raise DocumentProcessingException(f"Invalid DOCX file: {str(e)}")
    
    def parse(self) -> Dict[str, Any]:
        """
        Parse document and return structured JSON
        
        Returns:
            Dictionary containing document structure
        """
        try:
            parsed_data = {
                "metadata": self._extract_metadata(),
                "paragraphs": self._extract_paragraphs(),
                "total_paragraphs": len(self.document.paragraphs),
                "total_runs": self._count_total_runs()
            }
            
            logger.info(f"Parsed {parsed_data['total_paragraphs']} paragraphs, {parsed_data['total_runs']} runs")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Parsing failed: {str(e)}")
            raise DocumentProcessingException(f"Failed to parse document: {str(e)}")
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract document metadata"""
        core_properties = self.document.core_properties
        
        return {
            "title": core_properties.title or "",
            "author": core_properties.author or "",
            "subject": core_properties.subject or "",
            "created": str(core_properties.created) if core_properties.created else None,
            "modified": str(core_properties.modified) if core_properties.modified else None,
            "revision": core_properties.revision
        }
    
    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """Extract all paragraphs with their formatting"""
        paragraphs = []
        
        for idx, paragraph in enumerate(self.document.paragraphs):
            para_data = {
                "id": idx,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "alignment": self._get_alignment(paragraph.alignment),
                "runs": self._extract_runs(paragraph),
                "is_empty": len(paragraph.text.strip()) == 0
            }
            paragraphs.append(para_data)
        
        return paragraphs
    
    def _extract_runs(self, paragraph) -> List[Dict[str, Any]]:
        """Extract runs (text fragments with formatting) from a paragraph"""
        runs = []
        
        for run_idx, run in enumerate(paragraph.runs):
            run_data = {
                "id": run_idx,
                "text": run.text,
                "bold": run.bold if run.bold is not None else False,
                "italic": run.italic if run.italic is not None else False,
                "underline": run.underline if run.underline is not None else False,
                "font_name": run.font.name if run.font.name else "Calibri",
                "font_size": run.font.size.pt if run.font.size else 11,
                "font_color": self._get_color(run.font.color),
                "highlight": run.font.highlight_color if hasattr(run.font, 'highlight_color') else None
            }
            runs.append(run_data)
        
        return runs
    
    def _get_alignment(self, alignment) -> str:
        """Convert alignment enum to string"""
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            None: "left"
        }
        return alignment_map.get(alignment, "left")
    
    def _get_color(self, color) -> Optional[str]:
        """Extract RGB color as hex string"""
        if color and color.rgb:
            try:
                rgb = color.rgb
                return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
            except:
                return None
        return None
    
    def _count_total_runs(self) -> int:
        """Count total number of runs in document"""
        total = 0
        for paragraph in self.document.paragraphs:
            total += len(paragraph.runs)
        return total
    
    def get_paragraph_by_id(self, paragraph_id: int) -> Optional[Dict[str, Any]]:
        """
        Get specific paragraph by ID
        
        Args:
            paragraph_id: Paragraph index
            
        Returns:
            Paragraph data or None if not found
        """
        parsed = self.parse()
        paragraphs = parsed.get("paragraphs", [])
        
        for para in paragraphs:
            if para["id"] == paragraph_id:
                return para
        
        return None
    
    def get_text_content(self) -> str:
        """
        Extract plain text content from document
        
        Returns:
            Full document text
        """
        return "\n".join([para.text for para in self.document.paragraphs])
    
    def get_paragraph_texts(self) -> List[str]:
        """
        Get list of paragraph texts
        
        Returns:
            List of paragraph text strings
        """
        return [para.text for para in self.document.paragraphs]
