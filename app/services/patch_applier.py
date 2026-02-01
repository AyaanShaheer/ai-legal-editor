from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from io import BytesIO
from typing import List, Dict, Any, Optional
from datetime import datetime
from app.core.logging import logger
from app.core.exceptions import DocumentProcessingException


class PatchApplier:
    """Apply patches to DOCX documents with tracked changes"""
    
    def __init__(self, document_bytes: bytes):
        """
        Initialize applier with document content
        
        Args:
            document_bytes: Original DOCX file bytes
        """
        try:
            self.document = DocxDocument(BytesIO(document_bytes))
            logger.info("Document loaded for patch application")
        except Exception as e:
            logger.error(f"Failed to load document: {str(e)}")
            raise DocumentProcessingException(f"Invalid document: {str(e)}")
    
    def apply_paragraph_patches(
        self, 
        patches: List[Dict[str, Any]],
        author: str = "AI Assistant"
    ) -> bytes:
        """
        Apply multiple paragraph patches with tracked changes
        
        Args:
            patches: List of paragraph patch dictionaries
            author: Author name for tracked changes
            
        Returns:
            Modified document as bytes
        """
        try:
            applied_count = 0
            
            for patch in patches:
                paragraph_id = patch.get("paragraph_id")
                replacement_text = patch.get("replacement_text", "")
                
                # Apply patch to specific paragraph
                success = self._apply_single_paragraph_patch(
                    paragraph_id, 
                    replacement_text, 
                    author
                )
                
                if success:
                    applied_count += 1
            
            logger.info(f"Applied {applied_count}/{len(patches)} patches successfully")
            
            # Save to bytes
            output = BytesIO()
            self.document.save(output)
            output.seek(0)
            
            return output.read()
            
        except Exception as e:
            logger.error(f"Failed to apply patches: {str(e)}")
            raise DocumentProcessingException(f"Patch application failed: {str(e)}")
    
    def _apply_single_paragraph_patch(
        self, 
        paragraph_id: int, 
        new_text: str,
        author: str
    ) -> bool:
        """
        Apply patch to a single paragraph with tracked changes
        
        Args:
            paragraph_id: Paragraph index
            new_text: Replacement text
            author: Author name
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Get paragraph
            if paragraph_id >= len(self.document.paragraphs):
                logger.error(f"Paragraph {paragraph_id} not found")
                return False
            
            paragraph = self.document.paragraphs[paragraph_id]
            original_text = paragraph.text
            
            # Get formatting from first run (if exists)
            original_formatting = None
            if paragraph.runs:
                original_formatting = self._extract_run_formatting(paragraph.runs[0])
            
            # Clear all runs
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Add deletion markup for original text
            if original_text:
                self._add_deleted_text(paragraph, original_text, author, original_formatting)
            
            # Add insertion markup for new text
            if new_text:
                self._add_inserted_text(paragraph, new_text, author, original_formatting)
            
            logger.info(f"Applied patch to paragraph {paragraph_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to apply paragraph patch: {str(e)}")
            return False
    
    def _extract_run_formatting(self, run) -> Dict[str, Any]:
        """Extract formatting properties from a run"""
        return {
            "bold": run.bold if run.bold is not None else False,
            "italic": run.italic if run.italic is not None else False,
            "underline": run.underline if run.underline is not None else False,
            "font_name": run.font.name,
            "font_size": run.font.size
        }
    
    def _apply_formatting(self, run, formatting: Optional[Dict[str, Any]]):
        """Apply formatting to a run"""
        if not formatting:
            return
        
        run.bold = formatting.get("bold", False)
        run.italic = formatting.get("italic", False)
        run.underline = formatting.get("underline", False)
        
        if formatting.get("font_name"):
            run.font.name = formatting["font_name"]
        
        if formatting.get("font_size"):
            run.font.size = formatting["font_size"]
    
    def _add_deleted_text(
        self, 
        paragraph, 
        text: str, 
        author: str,
        formatting: Optional[Dict[str, Any]] = None
    ):
        """
        Add deleted text with track changes markup
        
        Args:
            paragraph: Paragraph object
            text: Text to mark as deleted
            author: Author name
            formatting: Original formatting
        """
        # Create deletion run
        run = paragraph.add_run(text)
        
        # Apply original formatting
        self._apply_formatting(run, formatting)
        
        # Add deletion markup (strikethrough + color)
        run.font.strike = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for deletions
        
        logger.debug(f"Added deleted text: {text[:50]}...")
    
    def _add_inserted_text(
        self, 
        paragraph, 
        text: str, 
        author: str,
        formatting: Optional[Dict[str, Any]] = None
    ):
        """
        Add inserted text with track changes markup
        
        Args:
            paragraph: Paragraph object
            text: Text to mark as inserted
            author: Author name
            formatting: Original formatting
        """
        # Create insertion run
        run = paragraph.add_run(text)
        
        # Apply original formatting
        self._apply_formatting(run, formatting)
        
        # Add insertion markup (underline + color)
        run.font.underline = True
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color for insertions
        
        logger.debug(f"Added inserted text: {text[:50]}...")
    
    def apply_simple_replacement(
        self, 
        paragraph_id: int, 
        new_text: str
    ) -> bool:
        """
        Apply simple replacement without tracked changes (for testing)
        
        Args:
            paragraph_id: Paragraph index
            new_text: Replacement text
            
        Returns:
            True if successful
        """
        try:
            if paragraph_id >= len(self.document.paragraphs):
                return False
            
            paragraph = self.document.paragraphs[paragraph_id]
            
            # Clear and replace
            paragraph.clear()
            paragraph.add_run(new_text)
            
            return True
            
        except Exception as e:
            logger.error(f"Simple replacement failed: {str(e)}")
            return False
    
    def get_document_bytes(self) -> bytes:
        """
        Get modified document as bytes
        
        Returns:
            Document bytes
        """
        output = BytesIO()
        self.document.save(output)
        output.seek(0)
        return output.read()
    
    def save_to_file(self, output_path: str):
        """
        Save modified document to file
        
        Args:
            output_path: Output file path
        """
        self.document.save(output_path)
        logger.info(f"Document saved to: {output_path}")


class BatchPatchApplier:
    """Apply patches to multiple documents"""
    
    @staticmethod
    def apply_patches_to_document(
        document_bytes: bytes,
        patches: List[Dict[str, Any]],
        author: str = "AI Assistant"
    ) -> bytes:
        """
        Convenience method to apply patches
        
        Args:
            document_bytes: Original document bytes
            patches: List of patches
            author: Author name
            
        Returns:
            Modified document bytes
        """
        applier = PatchApplier(document_bytes)
        return applier.apply_paragraph_patches(patches, author)
